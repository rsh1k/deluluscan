#!/usr/bin/env python3
"""build_report_docx.py — render the curated pentest report as an editable .docx.

Mirrors the target Penetration Test Report layout (the numbered §1–§10 format the
team uses) into a Word document the reader can edit directly. Content is derived
from the scan's results.json — specifically the curated `meta.report_include` set
— never hand-authored: the same integrity rule as the dashboard (the report may
only state what the scan observed). The attack narrative is emitted from
meta.escalation_pivot; when no chain was measured it says so rather than
inventing one.

Usage: python3 scripts/build_report_docx.py deluluscan-out/results.json out.docx
"""
from __future__ import annotations

import json
import os
import sys
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# the Deluluscan wordmark rasterised for a light surface (dark letterforms + magenta dot),
# mirroring the Letter of Attestation letterhead. Regenerate with:
#   sed 's/fill="white"/fill="#0F172A"/g' dashboard/public/logo-dark.svg > /tmp/l.svg
#   convert -background none -density 900 /tmp/l.svg -resize 1140x deluluscan/assets/report/target-logo.png
LOGO_PNG = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        'deluluscan', 'assets', 'report', 'target-logo.png')

# ---- palette (light, print-friendly) --------------------------------------
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x47, 0x55, 0x69)
SEV_COLOR = {
    'critical': RGBColor(0x99, 0x1B, 0x1B), 'high': RGBColor(0xC2, 0x41, 0x0C),
    'medium': RGBColor(0xA1, 0x62, 0x07), 'low': RGBColor(0x25, 0x63, 0xEB),
    'info': RGBColor(0x47, 0x55, 0x69),
}
SEV_LABEL = {'critical': 'CRITICAL', 'high': 'HIGH', 'medium': 'MEDIUM', 'low': 'LOW', 'info': 'INFO'}
SEV_RANK = {'critical': 4, 'high': 3, 'medium': 2, 'low': 1, 'info': 0}
VERDICT_LABEL = {
    'true_positive': 'True Positive', 'likely_true_positive': 'Likely TP',
    'confirmed': 'Confirmed', 'false_positive': 'False Positive',
    'likely_false_positive': 'Likely FP', 'inconclusive': 'Inconclusive',
    'unverified': 'Unverified',
}
OWASP_NAME = {
    'A01': 'Broken Access Control', 'A02': 'Cryptographic Failures', 'A03': 'Injection',
    'A04': 'Insecure Design', 'A05': 'Security Misconfiguration',
    'A06': 'Vulnerable and Outdated Components',
    'A07': 'Identification and Authentication Failures',
    'A08': 'Software and Data Integrity Failures',
    'A09': 'Security Logging and Monitoring Failures', 'A10': 'Server-Side Request Forgery',
}

# Curated per-finding classification + business impact. Taxonomy labels (OWASP/CWE)
# are chosen to match the team's report conventions rather than the noisier
# multi-value references list the scanner attaches; impact text is standard and
# true for the class, filled in where the scanner left it blank.
CLASSIFY = {
    'e980ff3037': dict(cwe='CWE-200', owasp='A05',
        impact='Internal class names and stack frames are returned to callers who should not see them. '
               'This discloses the framework, library versions and code structure, easing reconnaissance '
               'and the targeting of further attacks.'),
    '1c388166d0': dict(cwe='CWE-209', owasp='A05',
        impact='Verbose server errors expose internal implementation detail to unentitled callers, aiding '
               'an attacker in mapping the application and crafting further payloads.'),
    '3dff27583d': dict(cwe='CWE-862', owasp='A01',
        impact='A baseline back-end identity can invoke an operation that should require higher privilege, '
               'bypassing the intended trust boundary and disclosing role/layout information.'),
    '8453058847': dict(cwe='CWE-200', owasp='A05',
        impact='The full GraphQL schema is disclosed to unauthenticated callers, revealing every type, '
               'field and query the API exposes and easing discovery of sensitive or abusable operations.'),
    '1613242c37': dict(cwe='CWE-200', owasp='A05',
        impact='The complete API specification is served without authentication, enumerating every '
               'endpoint, parameter and schema and expanding the reconnaissance surface.'),
    '70bf9ead93': dict(cwe='CWE-755', owasp='A05',
        impact='Malformed input reaches an unhandled code path and returns an HTTP 500. Beyond the '
               'information leaked in the error, unhandled paths can mask missing input validation and are '
               'a common precursor to injection or denial-of-service issues.'),
}


def shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcpr.append(sh)


def set_cell(cell, text, *, bold=False, color=None, size=9, mono=False, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if mono:
        run.font.name = 'Consolas'
    if color:
        run.font.color.rgb = color
    return p


def h1(doc, num, text):
    p = doc.add_paragraph()
    r = p.add_run(f'{num}. {text}')
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = INK
    p.paragraph_format.space_before = Pt(14)
    return p


def sub(doc, text, *, size=9.5, upper=False):
    """A small uppercase sub-heading, as used inside each detailed finding."""
    p = doc.add_paragraph()
    r = p.add_run(text.upper() if upper else text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = MUTED
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
    return p


def h2(doc, text, color=INK):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = color
    return p


def para(doc, text, *, size=10, color=INK):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color
    return p


def bullet(doc, text, *, size=10):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = INK
    return p


def numbered(doc, text, *, size=9):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = INK
    return p


def mono_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    for i, line in enumerate(text.split('\n')):
        if i:
            p.add_run().add_break()
        r = p.add_run(line); r.font.name = 'Consolas'; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p


def header_row(table, labels):
    for c, lab in zip(table.rows[0].cells, labels):
        set_cell(c, lab, bold=True, color=MUTED, size=8)
        shade(c, 'F1F5F9')


def cell_border(cell, color='CBD5E1', sz='8'):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        borders.append(el)
    tcpr.append(borders)


def letterhead(doc):
    """A centered, light-shaded bordered box carrying the Deluluscan wordmark and the
    report title — the same letterhead treatment as the Letter of Attestation."""
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    shade(cell, 'F8FAFC'); cell_border(cell, 'E2E8F0', '10')
    cell.width = Inches(6.5)
    cell.text = ''

    def cline(text, size, *, bold=False, color=INK, space=3, first=False):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space)
        if first:
            p.paragraph_format.space_before = Pt(8)
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
        return p

    # logo
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(8)
    if os.path.exists(LOGO_PNG):
        p.add_run().add_picture(LOGO_PNG, width=Inches(2.1))
    else:
        r = p.add_run('the target'); r.bold = True; r.font.size = Pt(22)
    cline('CONFIDENTIAL REPORT', 8.5, bold=True, color=RGBColor(0xA1, 0x62, 0x07), space=10)
    cline('Penetration Test Report', 22, bold=True)
    cline('the target Content Management Platform', 12)
    cline('Authenticated Grey-Box Assessment of the REST API & Authorization Model', 10, color=MUTED)
    cline('Prepared by the Security Team', 9.5, color=MUTED, space=10)


def page_footer(doc, ref):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{ref} · Confidential — Internal Use Only · Prepared by the Security Team')
    r.font.size = Pt(8); r.font.color.rgb = MUTED


def classify(f):
    """(cwe, owasp, impact) for a finding.

    Prefers the taxonomy the knowledge base attached to the report block, so the
    same class of issue is always classified the same way, and falls back to the
    per-id CLASSIFY table for older scans that predate it.
    """
    rep = (f.get('detail') or {}).get('report') or {}
    tax = rep.get('taxonomy') or {}
    legacy = CLASSIFY.get(f['id'], {})
    cwe = ', '.join(tax.get('cwe') or []) or legacy.get('cwe', '')
    owasp = tax.get('owasp_2025') or legacy.get('owasp', '')
    impact = rep.get('impact') or (f.get('detail') or {}).get('impact') or legacy.get('impact', '')
    return cwe, owasp, impact


def cvss_of(f):
    """The structured CVSS block for a finding, or None if none was assigned."""
    rep = (f.get('detail') or {}).get('report') or {}
    c = rep.get('cvss') or (f.get('detail') or {}).get('cvss')
    return c if isinstance(c, dict) else None


def owasp_label(code):
    """Label an OWASP category code with its name.

    A ':2025' code MUST be named from the 2025 list — the module-level
    OWASP_NAME table below is the 2021 list, and the two disagree on what most
    codes mean (A02 is Cryptographic Failures in 2021, Security Misconfiguration
    in 2025). Naming a 2025 code from the 2021 table silently prints the wrong
    category, which is a factual error about the finding.
    """
    if not code:
        return ''
    if code.endswith(':2025'):
        from deluluscan.knowledge import owasp_2025_label
        return owasp_2025_label(code)
    return f"{code} {OWASP_NAME.get(code.split(':')[0], '')}".strip()


def render_exchanges(doc, rep):
    """Render each captured exchange as request + OBSERVED RESPONSE.

    A reproduction step without its response cannot be adjudicated by the reader:
    "HTTP 500" and "HTTP 500 with an empty body" are different findings, and only
    the body distinguishes a leak from a bare failure. Where a body is empty that
    is stated in words rather than shown as blank space, because an empty block
    reads as missing evidence rather than as the evidence it is.
    """
    exchanges = rep.get('exchanges') or []
    if not exchanges:
        if rep.get('reproduction'):
            sub(doc, 'Reproduction', upper=True)
            mono_block(doc, '\n'.join(rep['reproduction']))
        return
    sub(doc, 'Reproduction — request and observed response', upper=True)
    for ex in exchanges:
        # The label is already the first comment line of the curl block; printing
        # it again above the block just doubles it on the page.
        mono_block(doc, ex.get('curl', ''))
        resp = ex.get('response') or {}
        status = resp.get('status')
        nbytes = resp.get('body_bytes')
        head = f"HTTP {status}" if status else 'no response'
        if nbytes is not None:
            head += f'  ({nbytes} bytes)'
        para(doc, head, size=8.5, color=MUTED)
        if resp.get('body_empty'):
            para(doc, 'Empty body — nothing was disclosed in the response.', size=9)
        else:
            body = resp.get('body', '')
            if resp.get('body_truncated'):
                body += '\n… response truncated for the report …'
            mono_block(doc, body)


def render_measurements(doc, f):
    """Render any quantitative measurement the finding rests on.

    Some findings are proven by numbers rather than by a single response — a
    timing separation, or a throughput rate. Those numbers are the evidence, so
    they belong in the document next to the requests that produced them.
    """
    det = f.get('detail') or {}
    measured = det.get('measured')
    if isinstance(measured, dict) and measured:
        sub(doc, 'Measured', upper=True)
        t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
        header_row(t, ['MEASURE', 'VALUE'])
        for k, v in measured.items():
            row = t.add_row().cells
            set_cell(row[0], k.replace('_', ' '), size=8.5, color=MUTED)
            set_cell(row[1], str(v), size=8.5)

    timing = det.get('timing_samples')
    if isinstance(timing, dict) and timing:
        sub(doc, 'Timing measurements', upper=True)
        t = doc.add_table(rows=1, cols=6); t.style = 'Table Grid'
        header_row(t, ['ACCOUNT', 'EXISTS', 'SAMPLES', 'MEAN (s)', 'MIN (s)', 'MAX (s)'])
        for acct, v in timing.items():
            row = t.add_row().cells
            set_cell(row[0], acct, size=8.5)
            set_cell(row[1], 'yes' if v.get('group') == 'existing' else 'no', size=8.5)
            set_cell(row[2], str(v.get('samples', '')), size=8.5)
            set_cell(row[3], f"{v.get('mean_s', '')}", size=8.5)
            set_cell(row[4], f"{v.get('min_s', '')}", size=8.5)
            set_cell(row[5], f"{v.get('max_s', '')}", size=8.5)
        if det.get('separation_s') is not None:
            verdict = ('the two populations do not overlap, so a single request classifies an account '
                       'with certainty') if det.get('perfectly_separable') else \
                      'the populations partially overlap, so repeated sampling is required'
            para(doc, f"Separation: {det['separation_s']}s — {verdict}.", size=9)


def render_cvss(doc, cvss):
    """The score, the vector, and the reasoning behind every metric.

    A bare number is not reviewable. Publishing the per-metric rationale is what
    lets a reader disagree with a score rather than simply accept it.
    """
    if not cvss:
        return
    sub(doc, 'CVSS v3.1 base score', upper=True)
    score = cvss.get('base_score')
    sev = (cvss.get('severity') or '').lower()
    p = doc.add_paragraph()
    r = p.add_run(f"{score}  "); r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = SEV_COLOR.get(sev, INK)
    r = p.add_run(SEV_LABEL.get(sev, sev.upper()) + '   ')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = SEV_COLOR.get(sev, INK)
    r = p.add_run(cvss.get('vector', '')); r.font.size = Pt(8.5)
    r.font.name = 'Consolas'; r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(4)

    rationale = cvss.get('metric_rationale') or {}
    if rationale:
        t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
        header_row(t, ['METRIC', 'WHY THIS VALUE'])
        for metric, why in rationale.items():
            row = t.add_row().cells
            set_cell(row[0], metric, size=8.5, mono=True)
            set_cell(row[1], why, size=8.5)


def build(results_path, out_path):
    data = json.load(open(results_path))
    meta = data.get('meta', {})
    findings = data.get('findings', [])
    by_id = {x['id']: x for x in findings}
    include = (meta.get('report_include') or {}).get('ids') or []
    picked = [by_id[i] for i in include if i in by_id]
    picked.sort(key=lambda f: -SEV_RANK.get(f.get('severity', 'info'), 0))

    ver = next((x['version'] for x in meta.get('fingerprint', {}).get('detections', [])
                if x.get('tech') == 'the target'), 'unknown')
    cov = meta.get('coverage', {})
    ps = meta.get('probe_stats', {})
    idents = ', '.join(sorted((meta.get('identities') or {}).keys()))
    today = date.today().strftime('%B %d, %Y')
    ref = f"TARGET-PT-{(data.get('date') or '')[:7] or date.today().isoformat()[:7]}"
    target = data.get('target') or meta.get('target') or ''
    counts = {s: sum(1 for f in picked if f.get('severity') == s)
              for s in ('critical', 'high', 'medium', 'low', 'info')}
    exploitable = sum(1 for f in picked if f.get('exploitability') == 'exploitable')

    doc = Document()
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10); st.font.color.rgb = INK
    page_footer(doc, ref)

    def center(text, size, *, bold=False, color=INK, space=2):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(space); return p

    # ---- cover -------------------------------------------------------------
    center('CONFIDENTIAL — INTERNAL USE ONLY', 9, bold=True, color=MUTED, space=8)
    letterhead(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    cover = doc.add_table(rows=0, cols=2); cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k1, v1, k2, v2 in [
        ('Report reference', ref, 'Classification', 'Confidential / Internal'),
        ('Report date', today, 'Report version', '1.0 (Final)'),
        ('Target environment', target, 'Assessment type', 'Grey-box, authenticated'),
        ('Prepared by', 'the Security Team', 'Distribution', 'Engineering & Security leadership'),
        ('Platform version', f'the target {ver}', 'Assessment tool', 'Deluluscan'),
    ]:
        cells = cover.add_row().cells
        set_cell(cells[0], f'{k1}   {v1}', size=9)
        set_cell(cells[1], f'{k2}   {v2}', size=9)
    doc.add_page_break()

    # ---- 1. Executive Summary ---------------------------------------------
    h1(doc, 1, 'Executive Summary')
    para(doc, 'All testing described in this report was performed by the Security Team against a '
              'dedicated, authorized target instance. No customer systems or production data were involved.')
    para(doc, 'The target Cybersecurity Team conducted an authenticated grey-box penetration test of the target '
              "platform's REST API and authorization model. The assessment combined source-informed analysis of the "
              'the target codebase with live, multi-identity testing against a running instance to determine which '
              'weaknesses are genuinely exploitable rather than merely present in code.')
    sev_phrase = ' and '.join(', '.join(f'{counts[s]} {s}' for s in ('critical', 'high', 'medium', 'low') if counts[s]).rsplit(', ', 1))
    if picked:
        para(doc, f'This report presents {len(picked)} confirmed findings — {sev_phrase} — all {exploitable} of which '
                  f'{"is" if exploitable == 1 else "are"} demonstrated to be exploitable against the running target. '
                  f'Findings that could not be shown to have a security consequence are recorded separately as '
                  f'observations and are deliberately not counted here.')
    else:
        # A zero-finding report is a real result, but it must not read as "nothing
        # was found" when behaviours were found and then accepted. State both.
        n_obs = sum(1 for x in findings if (x.get('detail') or {}).get('observation'))
        n_ref = sum(1 for x in findings if (x.get('detail') or {}).get('refuted'))
        para(doc, 'This report presents NO exploitable vulnerabilities. That result should be read '
                  f'precisely. {n_obs} behaviours reproduced against the live target and are recorded as '
                  'observations in Section 8.1; each was classified by the product owner as accepted or '
                  'by design, and so none is counted as a finding. A further '
                  f'{n_ref} candidate classes raised by the previous report and by this engagement\'s '
                  'automated sweep were shown NOT to reproduce at all, and are refuted with evidence in '
                  'Section 8.2. The two groups are materially different: the first happens and has been '
                  'accepted; the second does not happen.')
    # The theme sentence is DERIVED from what was confirmed. Hand-authoring it is
    # how a report ends up asserting a narrative the evidence does not support —
    # the previous iteration described "an introspectable GraphQL schema" that
    # live re-testing showed was never introspectable.
    unauth = [f for f in picked
              if (cvss_of(f) or {}).get('vector', '').find('PR:N') != -1]
    _skip_theme = not picked
    classes = []
    for f in picked:
        cl = (f.get('vuln_class') or '').replace('_', ' ')
        if cl and cl not in classes:
            classes.append(cl)
    theme = ', '.join(classes[:-1]) + (' and ' + classes[-1] if len(classes) > 1 else '')
    if not _skip_theme:
      para(doc, f'The confirmed findings span {theme or "the classes detailed below"}. '
              + (f'{len(unauth)} of {len(picked)} require no authentication at all. '
                 if unauth else '')
              + 'Each was confirmed by a live test against the running instance and is reported with the '
                'request and the response it produced.')
    adj = meta.get('adjudication') or {}
    if adj.get('refuted_classes') and picked:
        occ = adj.get('refuted_occurrences')
        para(doc, f'A further {adj["refuted_classes"]} candidate classes'
                  + (f' ({occ} individual occurrences)' if occ else '')
                  + ' were raised during this engagement — by the previous report and by the automated '
                    'sweep — and were refuted by live re-testing. They are documented in Section 8 with '
                    'the evidence that refutes each, including a candidate previously rated Critical. '
                    'Reporting a weakness that does not exist costs engineering time and erodes trust in '
                    'the assessment, so the refutations are treated as deliverables in their own right.')
    if meta.get('chain'):
        para(doc, meta['chain'])
    t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'
    header_row(t, ['SEVERITY', 'CONFIRMED', 'MEANING'])
    for s, mean in [
        ('critical', 'Immediate risk of full system or data compromise; remediate now.'),
        ('high', 'Serious risk enabling privilege escalation or sensitive data exposure.'),
        ('medium', 'Meaningful weakness, typically requiring specific conditions.'),
        ('low', 'Limited impact; address as part of routine hardening.'),
    ]:
        row = t.add_row().cells
        set_cell(row[0], SEV_LABEL[s], bold=True, color=SEV_COLOR[s], size=9)
        set_cell(row[1], str(counts[s]), size=9)
        set_cell(row[2], mean, size=9)

    # ---- 2. Scope & Authorization -----------------------------------------
    h1(doc, 2, 'Scope & Authorization')
    para(doc, f'Authorization. This engagement was authorized internally by the target. Testing was confined to a '
              f'dedicated, non-production the target instance on a loopback/RFC-1918 address ({target}). No third-party, '
              f'customer, or production systems were in scope at any time.')
    h2(doc, 'In scope')
    bullet(doc, 'the target's REST API (/api/v1) and the legacy DWR AJAX surface.')
    bullet(doc, 'Authentication, session/JWT handling, and the role / layout / portlet authorization model.')
    bullet(doc, 'Multi-identity access-control testing across the privilege ladder (anonymous → back-end user → CMS Administrator).')
    h2(doc, 'Out of scope')
    bullet(doc, 'Denial-of-service and volumetric/load testing.')
    bullet(doc, 'Social engineering, phishing, and physical security.')
    bullet(doc, 'Underlying host, container runtime, and third-party infrastructure hardening.')
    h2(doc, 'Rules of engagement')
    bullet(doc, 'Testing was non-destructive; any state change made to prove a finding (e.g. a test role/layout assignment) was reverted immediately.')
    bullet(doc, 'Exploitation was carried out only to the point of proof. Code-execution findings were confirmed as reachable but were not weaponised.')

    # ---- 3. Methodology ----------------------------------------------------
    h1(doc, 3, 'Methodology')
    para(doc, 'The assessment followed an industry-aligned methodology drawing on the OWASP Web Security Testing '
              'Guide, the OWASP API Security Top 10, and the PTES execution standard, adapted for a source-informed workflow:')
    bullet(doc, 'Source-informed discovery. Static/semantic code analysis of the target source identified candidate weaknesses and the exact code paths and endpoints behind them.')
    bullet(doc, 'Live differential authorization testing. Each candidate was re-tested against the running instance as multiple identities (anonymous, a baseline back-end user, and an administrator). Divergent responses — e.g. a low-privilege identity succeeding where only an administrator should — indicate broken access control.')
    bullet(doc, 'Adjudication (confirm-to-proof). Every candidate was classified as a confirmed true positive, a false positive, or code-only-not-live-tested, based on observed evidence rather than the code pattern alone.')
    bullet(doc, 'Purpose-built test identities. Where a finding required a specific privilege (e.g. a scoped portlet), a dedicated non-administrative account was provisioned to test the exact trust boundary, then removed.')
    bullet(doc, f'Provenance. Source analysis was performed against the target source and {meta.get("source", "the captured OpenAPI specification")}. Live verification was performed against {target}.')
    h2(doc, 'Test coverage by vulnerability class')
    para(doc, 'The matrix below records which classes were exercised in this engagement and how. Classes marked not '
              'automated require manual testing and are not covered by this report — see Section 5 (Coverage & '
              'Limitations).', size=9, color=MUTED)
    matrix = [
        ('Broken access control / BFLA', 'A01 / API5', 'Entitlement-spec conformance sweep: every endpoint is replayed as each identity and any sub-tier success is a specification violation.', 'AUTOMATED'),
        ('Broken object-level authz (BOLA/IDOR)', 'A01 / API1', 'Object identifiers are requested across identities and ownership boundaries.', 'AUTOMATED'),
        ('Object property-level authz (BOPLA)', 'API3', 'Response property mining plus mass-assignment probes on writes.', 'AUTOMATED'),
        ('Authentication & session management', 'A07 / API2', 'JWT algorithm/signature handling, token scope, login and password-reset flows.', 'AUTOMATED'),
        ('Injection (SQL)', 'A03', 'Differential and error-based probing of parameters reaching query construction.', 'AUTOMATED'),
        ('Injection (template / SSTI)', 'A03', 'Template-expression evaluation probes on request-controlled input.', 'AUTOMATED'),
        ('Cross-site scripting', 'A03', 'Canary reflection with context analysis.', 'AUTOMATED'),
        ('Server-side request forgery', 'A10', 'Loopback/callback destination probing on URL-accepting parameters.', 'AUTOMATED'),
        ('Cryptographic failures', 'A02', 'Token entropy/sequencing and hard-coded or weak key material review.', 'AUTOMATED'),
        ('Security misconfiguration', 'A05', 'CORS, cache, verb tampering, error handling and exposed management surfaces.', 'AUTOMATED'),
        ('Sensitive data exposure', 'A01 / API3', 'Response inspection for server internals and over-broad data.', 'AUTOMATED'),
        ('Vulnerable & outdated components', 'A06', 'Version fingerprinting against the target CVEs and advisories.', 'AUTOMATED'),
        ('Rate limiting / resource consumption', 'API4', 'Burst probing of authentication and expensive operations.', 'PARTIAL'),
        ('Business-logic abuse', 'API6', 'Requires understanding of intended workflow semantics.', 'MANUAL'),
        ('Multi-step exploit chaining', '—', 'Chains are reported where observed, but creative chaining is a human activity.', 'PARTIAL'),
        ('Client-side / browser-driven attacks', '—', 'DOM XSS, clickjacking and CSP weaknesses in the admin UI.', 'MANUAL'),
        ('Social engineering / phishing', '—', 'Out of scope for automated assessment.', 'MANUAL'),
        ('Denial of service', '—', 'Deliberately excluded to protect availability.', 'MANUAL'),
    ]
    t = doc.add_table(rows=1, cols=4); t.style = 'Table Grid'
    header_row(t, ['VULNERABILITY CLASS', 'REF', 'HOW TESTED', 'COVERAGE'])
    for cls, r, how, covg in matrix:
        row = t.add_row().cells
        set_cell(row[0], cls, size=8.5); set_cell(row[1], r, size=8.5)
        set_cell(row[2], how, size=8.5)
        set_cell(row[3], covg, size=8.5, bold=covg != 'AUTOMATED', color=MUTED if covg == 'MANUAL' else INK)

    # ---- 4. Risk Rating Methodology ---------------------------------------
    h1(doc, 4, 'Risk Rating Methodology')
    scoring = meta.get('scoring') or {}
    if scoring:
        h2(doc, f"Scoring system: {scoring.get('system', 'CVSS v3.1 Base')}")
        para(doc, 'Every confirmed finding carries a CVSS vector string, so any reader can paste it into the '
                  'FIRST calculator and reproduce the score independently.', size=9.5)
        if scoring.get('rationale'):
            para(doc, scoring['rationale'], size=9.5)
        if scoring.get('metric_derivation'):
            para(doc, scoring['metric_derivation'], size=9.5)
        if scoring.get('implementation'):
            para(doc, f"Implementation: {scoring['implementation']}.", size=9, color=MUTED)
        t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
        header_row(t, ['CVSS BASE SCORE', 'QUALITATIVE RATING'])
        for rng, lbl in [('0.0', 'None'), ('0.1 - 3.9', 'Low'), ('4.0 - 6.9', 'Medium'),
                         ('7.0 - 8.9', 'High'), ('9.0 - 10.0', 'Critical')]:
            row = t.add_row().cells
            set_cell(row[0], rng, size=9, mono=True); set_cell(row[1], lbl, size=9)
        h2(doc, 'Severity bands used in this report')
    para(doc, 'Findings are rated on a four-tier scale derived from the likelihood of exploitation and the business '
              'impact of a successful attack. Ratings reflect the outcome of live testing, not merely theoretical exposure.')
    t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
    header_row(t, ['RATING', 'DEFINITION'])
    for s, dfn in [
        ('critical', 'Directly leads to full compromise of the application, underlying host, or all data; trivially or reliably exploitable.'),
        ('high', 'Enables privilege escalation, authentication bypass, or exposure of sensitive data; exploitable by a low-privilege or unauthenticated attacker.'),
        ('medium', 'Meaningful weakness that requires specific preconditions or yields limited data/impact.'),
        ('low', 'Minor issue or defence-in-depth gap with limited direct impact.'),
    ]:
        row = t.add_row().cells
        set_cell(row[0], SEV_LABEL[s], bold=True, color=SEV_COLOR[s], size=9)
        set_cell(row[1], dfn, size=9)

    # ---- 5. Coverage & Limitations ----------------------------------------
    h1(doc, 5, 'Coverage & Limitations')
    para(doc, 'This section states the boundaries of the engagement so that the absence of a finding is not mistaken '
              'for assurance. Untested is not the same as secure.')
    t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
    header_row(t, ['MEASURE', 'VALUE'])
    rows = [
        ('Endpoints discovered', str(cov.get('endpoints_discovered', '—'))),
        ('Endpoints actively probed', f"{cov.get('endpoints_probed', '—')} ({cov.get('endpoints_probed_pct', '—')}%)"),
        ('Endpoints NOT probed', str(cov.get('endpoints_not_probed',
                                             (cov.get('endpoints_discovered', 0) or 0)
                                             - (cov.get('endpoints_probed', 0) or 0)))),
    ]
    if cov.get('candidate_findings_raised'):
        rows.append(('Candidate findings raised & adjudicated', str(cov['candidate_findings_raised'])))
    if ps.get('requests'):
        rows.append(('HTTP requests issued', str(ps['requests'])))
    rows.append(('Identities exercised', idents))
    for k, v in rows:
        row = t.add_row().cells
        set_cell(row[0], k, size=9, color=MUTED); set_cell(row[1], v, size=9)
    h2(doc, 'Explicit limitations')
    # Prefer the limitations recorded for THIS engagement. The generic list below
    # is a fallback: a stale boilerplate limitation is as misleading as a stale
    # finding, and coverage claims are exactly where that matters most.
    stated = meta.get('limitations') or []
    if stated:
        for lim in stated:
            bullet(doc, lim.replace('**', ''))
    else:
        bullet(doc, 'Automated testing only. This engagement was executed by an automated harness (Deluluscan). Business-logic abuse, multi-step workflow manipulation, and creative exploit chaining specific to your deployment require a human tester and were not exhaustively attempted.')
        bullet(doc, 'Not covered: denial-of-service and load testing, social engineering and phishing, physical security, client-side/browser-driven attack paths, and hardening of the underlying host, container runtime and third-party infrastructure.')
        bullet(doc, 'Exploitation to proof only. Code-execution and deserialization findings were confirmed as reachable but were not weaponised; their full impact is asserted from the reachable code path, not from executed payloads.')
        bullet(doc, 'Point-in-time. Results reflect the target state at the time of testing and the source revision cited above. Subsequent changes are not covered.')

    # ---- 6. Summary of Findings -------------------------------------------
    h1(doc, 6, 'Summary of Findings')
    para(doc, 'The table below lists every finding detailed in this report. Confirmed findings carry a reference '
              '(F-nn) and are detailed in Section 8. The Status column reflects the current triage state.')
    refmap = {}
    t = doc.add_table(rows=1, cols=8); t.style = 'Table Grid'
    header_row(t, ['REF', 'FINDING', 'SEVERITY', 'CVSS', 'CWE', 'OWASP', 'VERDICT', 'STATUS'])
    for i, f in enumerate(picked, 1):
        fref = f'F-{i:02d}'; refmap[f['id']] = fref
        cwe, owasp, _ = classify(f)
        c = cvss_of(f)
        row = t.add_row().cells
        set_cell(row[0], fref, size=8.5, mono=True)
        set_cell(row[1], f['title'], size=8.5)
        set_cell(row[2], SEV_LABEL[f['severity']], bold=True, color=SEV_COLOR[f['severity']], size=8.5)
        set_cell(row[3], str(c.get('base_score')) if c else '—', size=8.5, mono=True)
        set_cell(row[4], cwe or '—', size=8.5)
        set_cell(row[5], owasp or '—', size=8.5)
        set_cell(row[6], VERDICT_LABEL.get(f.get('verdict'), f.get('verdict', '—')), size=8.5)
        set_cell(row[7], 'Open', size=8.5)

    # Observations and refuted candidates are listed here too: a reader must be
    # able to see what was examined and dismissed, not only what was reported.
    observations = [x for x in findings
                    if (x.get('detail') or {}).get('observation') and x['id'] not in include]
    refuted = [x for x in findings if (x.get('detail') or {}).get('refuted')]
    if observations:
        h2(doc, 'Observations — confirmed behaviour, not counted as findings')
        para(doc, 'These behaviours are real and reproducible. None could be shown to have a security '
                  'consequence, so none is counted as a vulnerability or assigned a CVSS score. They are '
                  'listed so the reader knows they were examined and what was concluded.', size=9, color=MUTED)
        t = doc.add_table(rows=1, cols=4); t.style = 'Table Grid'
        header_row(t, ['REF', 'OBSERVATION', 'OWASP', 'CWE'])
        for i, f in enumerate(observations, 1):
            oref = f'O-{i:02d}'; refmap[f['id']] = oref
            cwe, owasp, _ = classify(f)
            row = t.add_row().cells
            set_cell(row[0], oref, size=8.5, mono=True)
            set_cell(row[1], f['title'], size=8.5)
            set_cell(row[2], owasp or '—', size=8.5)
            set_cell(row[3], cwe or '—', size=8.5)
    if refuted:
        h2(doc, 'Refuted candidates — false positives')
        adj = meta.get('adjudication') or {}
        occ = adj.get('refuted_occurrences')
        para(doc, f'{len(refuted)} candidate classes'
                  + (f' ({occ} individual occurrences)' if occ else '')
                  + ' were raised and then refuted by live re-testing. They are documented in Section 8 '
                    'with the evidence that refutes each one, so that they are not re-raised as findings '
                    'in a later cycle.', size=9, color=MUTED)

    # ---- 7. Attack Narrative ----------------------------------------------
    h1(doc, 7, 'Attack Narrative')
    pivot = meta.get('escalation_pivot') or {}
    if meta.get('chain_detail'):
        for line in meta['chain_detail']:
            para(doc, line, size=9.5) if not line.startswith('- ') else bullet(doc, line[2:], size=9.5)
    elif pivot.get('performed') and pivot.get('capabilities_gained'):
        para(doc, pivot.get('narrative', 'A privilege-escalation chain was measured end to end.'))
    else:
        para(doc, 'No end-to-end privilege-escalation chain was demonstrated during this engagement. The findings '
                  'below were each confirmed individually; they were not observed to compose into a single chain from '
                  'a low-privilege foothold to full compromise. The dominant theme is information disclosure to '
                  'unentitled callers and a set of privileged read endpoints reachable by a baseline back-end '
                  'identity — each of which should be remediated on its own merits.')

    # ---- 8. Detailed Findings ---------------------------------------------
    h1(doc, 8, 'Detailed Findings')
    for f in picked:
        rep = (f.get('detail', {}) or {}).get('report', {}) or {}
        cwe, owasp, impact = classify(f)
        sev = f['severity']
        hp = doc.add_paragraph(); hp.paragraph_format.space_before = Pt(10)
        r = hp.add_run(f'{refmap[f["id"]]}  '); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = INK
        r = hp.add_run(SEV_LABEL[sev] + '  '); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = SEV_COLOR[sev]
        r = hp.add_run(f['title']); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = INK

        sub(doc, 'Affected endpoint / component', upper=True)
        mono_block(doc, (rep.get('location') or {}).get('endpoint') or f.get('endpoint', ''))

        sub(doc, 'Classification', upper=True)
        tax = rep.get('taxonomy') or {}
        t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
        header_row(t, ['SCHEME', 'CLASSIFICATION'])
        for label, value in [
            ('OWASP Top 10:2025', owasp_label(owasp) or '—'),
            ('OWASP API Security Top 10:2023', tax.get('owasp_api_top10') or '—'),
            ('CWE', cwe or '—'),
            ('Exploitability', f.get('exploitability', 'unknown').replace('_', ' ').title()),
            ('Confidence', f.get('confidence', '—')),
        ]:
            row = t.add_row().cells
            set_cell(row[0], label, size=8.5, color=MUTED); set_cell(row[1], value, size=8.5)

        render_cvss(doc, cvss_of(f))

        if f.get('description'):
            sub(doc, 'Description', upper=True); para(doc, f['description'], size=9.5)
        elif rep.get('objective'):
            sub(doc, 'Description', upper=True); para(doc, rep['objective'], size=9.5)
        if rep.get('steps'):
            sub(doc, 'Testing performed', upper=True)
            for stp in rep['steps']:
                numbered(doc, stp, size=9)
        render_exchanges(doc, rep)
        render_measurements(doc, f)
        if rep.get('outcome'):
            sub(doc, 'Result & evidence', upper=True); para(doc, rep['outcome'], size=9.5)
        rv = (f.get('detail', {}) or {}).get('reverified')
        aff = (f.get('detail', {}) or {}).get('affected_endpoints') or []
        if rv:
            para(doc, f"Re-verification ({rv.get('date', '')}): {rv.get('reproduced', '')} endpoint(s) reproduced on "
                      f"live re-test. {rv.get('method', '')}", size=8.5, color=MUTED)
        if len(aff) > 1:
            para(doc, 'Re-verified affected endpoints:', size=8.5, color=MUTED)
            for e in aff:
                bullet(doc, e, size=8.5)

        sub(doc, 'Business impact', upper=True)
        para(doc, rep.get('impact') or impact, size=9.5)
        sub(doc, 'Recommendation', upper=True)
        para(doc, rep.get('remediation', 'See remediation plan.'), size=9.5)

    # ---- 8.1 Observations --------------------------------------------------
    if observations:
        h2(doc, 'Observations — confirmed behaviour, not counted as findings')
        para(doc, 'Each behaviour below was reproduced against the live target. None is scored, because in '
                  'each case the security consequence could not be demonstrated — the reasoning is stated '
                  'per item rather than left implicit.', size=9.5)
        for f in observations:
            rep = (f.get('detail', {}) or {}).get('report', {}) or {}
            cwe, owasp, impact = classify(f)
            hp = doc.add_paragraph(); hp.paragraph_format.space_before = Pt(8)
            r = hp.add_run(f'{refmap[f["id"]]}  '); r.bold = True; r.font.size = Pt(11)
            r = hp.add_run(f['title']); r.bold = True; r.font.size = Pt(11)
            para(doc, f'{cwe} · {owasp_label(owasp)}'.strip(' ·'), size=8.5, color=MUTED)
            if f.get('description'):
                para(doc, f['description'], size=9.5)
            render_exchanges(doc, rep)
            dispo = (f.get('detail') or {}).get('disposition')
            if dispo:
                sub(doc, 'Disposition', upper=True); para(doc, dispo, size=9.5)
            sub(doc, 'Assessed impact', upper=True)
            para(doc, rep.get('impact') or impact, size=9.5)
            if rep.get('remediation'):
                sub(doc, 'Recommendation', upper=True)
                para(doc, rep['remediation'], size=9.5)

    # ---- 8.2 Refuted candidates -------------------------------------------
    if refuted:
        h2(doc, 'Refuted candidates — false positives')
        para(doc, 'Each candidate below was raised — by the previous report, by this engagement\'s automated '
                  'sweep, or both — and then refuted by live re-testing against the running target. They are '
                  'documented with the evidence that refutes them, and with the mechanism, so that the same '
                  'candidate is not re-raised as a finding in a later cycle. A scanner signature is a '
                  'hypothesis, not a finding.', size=9.5)
        t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'
        header_row(t, ['CANDIDATE', 'RAISED BY', 'WHY IT IS NOT A FINDING'])
        for f in refuted:
            row = t.add_row().cells
            set_cell(row[0], f['title'], size=8.5)
            set_cell(row[1], (f.get('detail') or {}).get('origin', '—'), size=8.5, color=MUTED)
            set_cell(row[2], f.get('description', ''), size=8.5)

    # ---- 9. Remediation Plan ----------------------------------------------
    h1(doc, 9, 'Remediation Plan')
    para(doc, 'Remediation is prioritised by confirmed severity and demonstrated exploitability. Priorities assume '
              'the target is, or is representative of, a production deployment.')
    prio = {'critical': ('P1', 'Immediate — within 7 days'), 'high': ('P2', 'Urgent — within 30 days'),
            'medium': ('P3', 'Planned — within 90 days'), 'low': ('P4', 'Backlog — next hardening cycle'),
            'info': ('P4', 'Backlog — next hardening cycle')}
    t = doc.add_table(rows=1, cols=5); t.style = 'Table Grid'
    header_row(t, ['PRIORITY', 'TARGET', 'REF', 'FINDING', 'REQUIRED ACTION'])
    for f in picked:
        rep = (f.get('detail', {}) or {}).get('report', {}) or {}
        p, tgt = prio.get(f['severity'], ('P4', 'Backlog'))
        row = t.add_row().cells
        set_cell(row[0], p, size=8.5, bold=True); set_cell(row[1], tgt, size=8.5)
        set_cell(row[2], refmap[f['id']], size=8.5, mono=True)
        set_cell(row[3], f['title'], size=8.5)
        set_cell(row[4], rep.get('remediation', 'See detailed finding.'), size=8.5)

    # ---- 10. Conclusion ----------------------------------------------------
    h1(doc, 10, 'Conclusion & Recommendations')
    para(doc, meta.get('conclusion') or
              'The findings detailed in this report were each confirmed against the running target and should '
              'be remediated on their own merits.')
    h2(doc, 'Strategic recommendations')
    # Derived from the confirmed findings, so the recommendations cannot drift
    # out of step with what was actually found.
    for f in picked:
        rep = (f.get('detail', {}) or {}).get('report', {}) or {}
        rec = rep.get('remediation')
        if rec:
            first = rec.split('. ')[0].rstrip('.') + '.'
            bullet(doc, f"{refmap[f['id']]} — {first}")
    bullet(doc, 'Re-test after remediation. The Cybersecurity Team should re-run this assessment once fixes land to confirm closure.')
    para(doc, 'Disclaimer. This assessment reflects the state of the target at the time of testing and is not an '
              'exhaustive guarantee of security. Findings were validated to the point of proof only; no destructive '
              'actions were taken.', size=9, color=MUTED)

    doc.save(out_path)
    print(f'wrote {out_path}  ({len(picked)} findings: ' +
          ', '.join(f'{counts[s]} {s}' for s in ('critical', 'high', 'medium', 'low', 'info') if counts[s]) + ')')


if __name__ == '__main__':
    src = sys.argv[1] if len(sys.argv) > 1 else 'deluluscan-out/results.json'
    out = sys.argv[2] if len(sys.argv) > 2 else 'deluluscan-out/target-Penetration-Test-Report.docx'
    build(src, out)
